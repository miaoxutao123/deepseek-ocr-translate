import os
import shutil
import json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import Optional, AsyncGenerator
from pathlib import Path
import asyncio
import logging

from ..database import get_db
from ..models import User, History, TaskStatus, TaskType
from ..schemas import OCRResponse, OCRPageResult
from ..services import OCRService
from ..config import settings
from ..utils import EncryptionManager
from .auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ocr", tags=["ocr"])


async def process_ocr_background(history_id: int, user_id: int):
    """Background task to process OCR"""
    from ..database import SessionLocal

    logger.info(f"=== Background OCR task STARTING ===")
    logger.info(f"history_id={history_id}, user_id={user_id}")

    db = SessionLocal()
    try:
        logger.info(f"Database session created")

        # Get history and user
        history = db.query(History).filter(History.id == history_id).first()
        user = db.query(User).filter(User.id == user_id).first()

        if not history or not user:
            logger.error(f"History or user not found: history_id={history_id}, user_id={user_id}")
            return

        logger.info(f"Found history: {history.id}, file_path: {history.file_path}")
        logger.info(f"Found user: {user.id}, username: {user.username}")

        # Update status
        history.status = TaskStatus.PROCESSING
        db.commit()
        logger.info(f"Status updated to PROCESSING")

        # Get OCR service
        logger.info(f"Creating OCR service...")
        ocr_service = get_user_ocr_service(user)
        logger.info(f"OCR service created")

        file_path = history.file_path
        file_ext = Path(file_path).suffix.lower()
        logger.info(f"File extension: {file_ext}, processing...")

        ocr_results = []

        try:
            if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                # Single image
                logger.info(f"Processing single image: {file_path}")
                history.total_pages = 1
                history.current_page = 0
                history.progress_message = "正在处理图片..."
                db.commit()

                result = await ocr_service.ocr_single_image(file_path)
                ocr_results = [{
                    "page_number": 1,
                    "text": result["text"],
                    "confidence": result.get("confidence")
                }]

                history.current_page = 1
                history.progress_message = "处理完成"
                db.commit()

            elif file_ext == '.pdf':
                # PDF processing
                logger.info(f"Processing PDF: {file_path}")
                from pdf2image import convert_from_path
                from pypdf import PdfReader

                # Get page count
                pdf_reader = PdfReader(file_path)
                total_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {total_pages} pages")

                # Initialize progress
                history.total_pages = total_pages
                history.current_page = 0
                history.progress_message = f"开始处理 PDF，共 {total_pages} 页"
                db.commit()
                logger.info(f"进度初始化完成: 总页数={total_pages}")

                convert_kwargs = {
                    'dpi': 150,
                    'fmt': 'png',
                    'thread_count': 1,
                }
                if settings.POPPLER_PATH:
                    convert_kwargs['poppler_path'] = settings.POPPLER_PATH
                    logger.info(f"使用 Poppler 路径: {settings.POPPLER_PATH}")

                temp_dir = Path(file_path).parent / f"temp_{history_id}"
                temp_dir.mkdir(exist_ok=True)
                logger.info(f"临时目录创建: {temp_dir}")

                try:
                    for page_num in range(1, total_pages + 1):
                        logger.info("=" * 60)
                        logger.info(f"开始处理第 {page_num}/{total_pages} 页")
                        logger.info("=" * 60)

                        # Update progress
                        history.current_page = page_num
                        history.progress_message = f"正在处理第 {page_num} 页，共 {total_pages} 页"
                        db.commit()
                        logger.info(f"进度更新: {page_num}/{total_pages}")

                        # Convert single page
                        logger.info(f"正在转换 PDF 第 {page_num} 页为图片...")
                        import time
                        convert_start = time.time()

                        images = convert_from_path(
                            file_path,
                            first_page=page_num,
                            last_page=page_num,
                            **convert_kwargs
                        )

                        convert_time = time.time() - convert_start
                        logger.info(f"PDF 转换完成，耗时 {convert_time:.2f} 秒")

                        if not images:
                            logger.warning(f"第 {page_num} 页转换失败，跳过")
                            continue

                        image = images[0]
                        temp_image_path = temp_dir / f"page_{page_num}.png"

                        save_start = time.time()
                        image.save(temp_image_path, 'PNG', optimize=True)
                        save_time = time.time() - save_start

                        image_size = temp_image_path.stat().st_size
                        logger.info(f"图片保存完成: {temp_image_path}")
                        logger.info(f"图片大小: {image_size/1024:.2f} KB，保存耗时 {save_time:.2f} 秒")

                        # OCR the page
                        logger.info(f"开始 OCR 识别第 {page_num} 页...")
                        ocr_start = time.time()

                        result = await ocr_service.ocr_single_image(str(temp_image_path))

                        ocr_time = time.time() - ocr_start
                        logger.info(f"OCR 识别完成，耗时 {ocr_time:.2f} 秒")

                        ocr_results.append({
                            "page_number": page_num,
                            "text": result["text"],
                            "confidence": result.get("confidence")
                        })

                        logger.info(f"第 {page_num} 页处理完成，识别文本长度: {len(result['text'])} 字符")

                        # Clean up
                        temp_image_path.unlink()
                        logger.info(f"临时文件清理完成")
                        del image, images

                finally:
                    if temp_dir.exists():
                        shutil.rmtree(temp_dir, ignore_errors=True)

            elif file_ext in ['.txt', '.md']:
                # Text file
                logger.info(f"Processing text file: {file_path}")
                history.total_pages = 1
                history.current_page = 0
                history.progress_message = "正在读取文本文件..."
                db.commit()

                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                ocr_results = [{
                    "page_number": 1,
                    "text": text,
                    "confidence": 1.0
                }]

                history.current_page = 1
                history.progress_message = "处理完成"
                db.commit()

            elif file_ext in ['.docx', '.doc']:
                # Word document
                logger.info(f"Processing Word document: {file_path}")
                history.total_pages = 1
                history.current_page = 0
                history.progress_message = "正在读取 Word 文档..."
                db.commit()

                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                ocr_results = [{
                    "page_number": 1,
                    "text": text,
                    "confidence": 1.0
                }]

                history.current_page = 1
                history.progress_message = "处理完成"
                db.commit()

            # Save results
            history.ocr_result = json.dumps(ocr_results)
            history.status = TaskStatus.COMPLETED
            history.progress_message = "全部完成"
            from datetime import datetime
            history.completed_at = datetime.utcnow()
            db.commit()

            logger.info(f"Background OCR task completed for history_id={history_id}")

        except Exception as e:
            logger.error(f"Background OCR task failed: {str(e)}", exc_info=True)
            import traceback
            logger.error(f"Full traceback:\n{traceback.format_exc()}")
            history.status = TaskStatus.FAILED
            history.error_message = str(e)
            db.commit()

    except Exception as e:
        logger.error(f"Fatal error in background task: {str(e)}", exc_info=True)
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
    finally:
        db.close()
        logger.info(f"=== Background OCR task ended for history_id={history_id} ===")


def get_user_ocr_service(user: User) -> OCRService:
    """Create OCR service for user"""
    if not user.ocr_api_key:
        raise HTTPException(status_code=400, detail="OCR API key not configured")

    api_key = EncryptionManager.decrypt_api_key(
        user.ocr_api_key, user.id, settings.SECRET_KEY
    )
    api_base = user.ocr_api_base or "https://api.siliconflow.cn/v1"
    model = user.ocr_model or "deepseek-ai/deepseek-vl2"

    return OCRService(api_base, api_key, model)


async def save_upload_file(upload_file: UploadFile, user_id: int) -> str:
    """Save uploaded file and return path"""
    # Create user-specific upload directory
    user_upload_dir = Path(settings.UPLOAD_DIR) / str(user_id)
    user_upload_dir.mkdir(parents=True, exist_ok=True)

    # Generate unique filename
    timestamp = int(asyncio.get_event_loop().time() * 1000)
    file_ext = Path(upload_file.filename).suffix
    filename = f"{timestamp}{file_ext}"
    file_path = user_upload_dir / filename

    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(upload_file.file, buffer)

    return str(file_path)


async def ocr_progress_stream(
    history_id: int,
    file_path: str,
    ocr_service: OCRService,
    db: Session
) -> AsyncGenerator[str, None]:
    """Stream OCR progress using SSE"""
    try:
        history = db.query(History).filter(History.id == history_id).first()
        history.status = TaskStatus.PROCESSING
        db.commit()

        # Determine file type and process
        file_ext = Path(file_path).suffix.lower()

        if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
            # Single image
            yield f"data: {json.dumps({'type': 'progress', 'message': 'Processing image...', 'page': 1})}\n\n"

            result = await ocr_service.ocr_single_image(file_path)

            ocr_results = [{
                "page_number": 1,
                "text": result["text"],
                "confidence": result.get("confidence")
            }]

        elif file_ext == '.pdf':
            # PDF - convert to images and OCR each page
            yield f"data: {json.dumps({'type': 'progress', 'message': 'Converting PDF to images...'})}\n\n"

            try:
                from pdf2image import convert_from_path
                from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="pdf2image library not installed. Please run: pip install pdf2image"
                )

            # Convert PDF to images with lower DPI to reduce memory usage
            convert_kwargs = {
                'dpi': 150,  # Reduced from 200 to save memory
                'fmt': 'png',
                'thread_count': 1,  # Single thread to avoid memory issues
            }
            if settings.POPPLER_PATH:
                convert_kwargs['poppler_path'] = settings.POPPLER_PATH

            ocr_results = []
            temp_dir = Path(file_path).parent / f"temp_{history_id}"
            temp_dir.mkdir(exist_ok=True)

            try:
                # Process PDF page by page instead of loading all at once
                # First, get page count
                from pypdf import PdfReader
                pdf_reader = PdfReader(file_path)
                total_pages = len(pdf_reader.pages)

                yield f"data: {json.dumps({'type': 'progress', 'message': f'Found {total_pages} pages, starting OCR...'})}\n\n"

                # Process each page individually to save memory
                for page_num in range(1, total_pages + 1):
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'Converting page {page_num}/{total_pages}...', 'page': page_num})}\n\n"

                    # Convert single page at a time
                    images = convert_from_path(
                        file_path,
                        first_page=page_num,
                        last_page=page_num,
                        **convert_kwargs
                    )

                    if not images:
                        continue

                    image = images[0]

                    # Save temporary image
                    temp_image_path = temp_dir / f"page_{page_num}.png"
                    image.save(temp_image_path, 'PNG', optimize=True)

                    yield f"data: {json.dumps({'type': 'progress', 'message': f'OCR processing page {page_num}/{total_pages}...', 'page': page_num})}\n\n"

                    # OCR the page
                    result = await ocr_service.ocr_single_image(str(temp_image_path))

                    ocr_results.append({
                        "page_number": page_num,
                        "text": result["text"],
                        "confidence": result.get("confidence")
                    })

                    # Clean up temp image immediately
                    temp_image_path.unlink()

                    # Clear the image from memory
                    del image
                    del images

            except PDFInfoNotInstalledError:
                raise HTTPException(
                    status_code=500,
                    detail="Poppler not found. Please install poppler or set POPPLER_PATH in .env file. "
                           "See setup_poppler.bat for installation instructions."
                )
            except Exception as e:
                error_msg = str(e)
                logger.error(f"PDF processing error: {error_msg}")
                if 'poppler' in error_msg.lower():
                    raise HTTPException(
                        status_code=500,
                        detail="Poppler not found. Please install poppler or set POPPLER_PATH in .env file. "
                               "See setup_poppler.bat for installation instructions."
                    )
                raise HTTPException(
                    status_code=500,
                    detail=f"PDF processing failed: {error_msg}"
                )
            finally:
                # Clean up temp directory
                if temp_dir.exists():
                    shutil.rmtree(temp_dir, ignore_errors=True)

        elif file_ext in ['.txt', '.md']:
            # Text file - just read content
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            ocr_results = [{
                "page_number": 1,
                "text": text,
                "confidence": 1.0
            }]

        elif file_ext in ['.docx', '.doc']:
            # Word document
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            ocr_results = [{
                "page_number": 1,
                "text": text,
                "confidence": 1.0
            }]

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

        # Save results to database
        history.ocr_result = json.dumps(ocr_results)
        history.status = TaskStatus.COMPLETED
        history.completed_at = db.query(History).filter(History.id == history_id).first().created_at
        db.commit()

        # Send completion message
        yield f"data: {json.dumps({'type': 'complete', 'results': ocr_results})}\n\n"

    except Exception as e:
        history = db.query(History).filter(History.id == history_id).first()
        history.status = TaskStatus.FAILED
        history.error_message = str(e)
        db.commit()

        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"


@router.post("/upload", response_model=dict)
async def upload_file_for_ocr(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    source_language: Optional[str] = Form("auto"),
    auto_process: bool = Form(True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a file for OCR processing"""
    # Validate file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Seek back to start

    if file_size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"File too large. Max size: {settings.MAX_FILE_SIZE_MB}MB")

    # Save file
    file_path = await save_upload_file(file, current_user.id)

    # Create history entry
    history = History(
        user_id=current_user.id,
        task_type=TaskType.OCR,
        status=TaskStatus.PENDING,
        original_filename=file.filename,
        file_path=file_path,
        file_size=file_size,
        source_language=source_language
    )

    db.add(history)
    db.commit()
    db.refresh(history)

    # If auto_process is True, start background task
    if auto_process:
        logger.info(f"Starting background OCR task for history_id={history.id}")
        background_tasks.add_task(
            process_ocr_background,
            history.id,
            current_user.id
        )

    return {
        "message": "File uploaded successfully",
        "history_id": history.id,
        "filename": file.filename,
        "auto_process": auto_process
    }


@router.get("/process/{history_id}")
async def process_ocr_stream(
    history_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Process OCR with SSE progress updates"""
    # Get history entry
    history = db.query(History).filter(
        History.id == history_id,
        History.user_id == current_user.id
    ).first()

    if not history:
        raise HTTPException(status_code=404, detail="History not found")

    if history.status == TaskStatus.COMPLETED:
        # Already processed, return cached result
        return {"status": "completed", "results": json.loads(history.ocr_result)}

    # Get OCR service
    ocr_service = get_user_ocr_service(current_user)

    # Return SSE stream
    return StreamingResponse(
        ocr_progress_stream(history_id, history.file_path, ocr_service, db),
        media_type="text/event-stream"
    )


@router.get("/result/{history_id}", response_model=OCRResponse)
def get_ocr_result(
    history_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get OCR results for a history entry"""
    history = db.query(History).filter(
        History.id == history_id,
        History.user_id == current_user.id
    ).first()

    if not history:
        raise HTTPException(status_code=404, detail="History not found")

    if history.status != TaskStatus.COMPLETED:
        raise HTTPException(status_code=400, detail=f"OCR not completed. Status: {history.status}")

    if not history.ocr_result:
        raise HTTPException(status_code=404, detail="OCR result not found")

    ocr_results = json.loads(history.ocr_result)
    pages = [OCRPageResult(**page) for page in ocr_results]

    return OCRResponse(
        task_id=history.id,
        pages=pages,
        total_pages=len(pages)
    )


@router.get("/status/{history_id}")
def get_ocr_status(
    history_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get OCR processing status"""
    history = db.query(History).filter(
        History.id == history_id,
        History.user_id == current_user.id
    ).first()

    if not history:
        raise HTTPException(status_code=404, detail="History not found")

    response = {
        "history_id": history.id,
        "status": history.status,
        "filename": history.original_filename,
        "created_at": history.created_at.isoformat() if history.created_at else None,
        "completed_at": history.completed_at.isoformat() if history.completed_at else None,
    }

    if history.status == TaskStatus.COMPLETED and history.ocr_result:
        ocr_results = json.loads(history.ocr_result)
        response["total_pages"] = len(ocr_results)
        response["has_result"] = True
    elif history.status == TaskStatus.FAILED:
        response["error_message"] = history.error_message
        response["has_result"] = False
    else:
        response["has_result"] = False

    return response
